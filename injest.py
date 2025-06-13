from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import os

DOCS_DIR = "docs"
IMG_DIR = "static/images"
os.makedirs(IMG_DIR, exist_ok=True)

embed_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
documents = []

for fname in os.listdir(DOCS_DIR):
    if fname.endswith(".docx"):
        fpath = os.path.join(DOCS_DIR, fname)
        doc = Document(fpath)
        # Extract text
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        # Extract tables as HTML
        tables = []
        for table in doc.tables:
            table_html = "<table border='1'>"
            for row in table.rows:
                table_html += "<tr>" + "".join(f"<td>{cell.text}</td>" for cell in row.cells) + "</tr>"
            table_html += "</table>"
            tables.append(table_html)
        # Extract images
        image_paths = []
        rels = doc.part.rels
        for rel in rels:
            rel = rels[rel]
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                img_name = f"{fname}_{rel.target_ref.split('/')[-1]}"
                img_path = os.path.join(IMG_DIR, img_name)
                with open(img_path, "wb") as f:
                    f.write(img_data)
                image_paths.append(img_path.replace("\\", "/"))
        # Store as one document per file
        content = full_text
        if tables:
            content += "\n\n[TABLES]\n" + "\n".join(tables)
        if image_paths:
            content += "\n\n[IMAGES]\n" + "\n".join(image_paths)
        documents.append({"page_content": content, "metadata": {"source": fname}})

# Split and embed
splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=400)
docs = splitter.create_documents([d["page_content"] for d in documents], metadatas=[d["metadata"] for d in documents])
db = FAISS.from_documents(docs, embed_model)
db.save_local("vectorstore")
print("Vectorstore built!")